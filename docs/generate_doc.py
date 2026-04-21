from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# Sayfa kenar boşlukları
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2.5)

DARK_BLUE = RGBColor(0x1F, 0x38, 0x64)
MED_BLUE  = RGBColor(0x2E, 0x54, 0x96)
GREEN     = RGBColor(0x1F, 0x6B, 0x00)
GRAY_TXT  = RGBColor(0x55, 0x55, 0x55)
BLACK     = RGBColor(0x1A, 0x1A, 0x1A)
YELLOW_BG = RGBColor(0xFF, 0xF2, 0xCC)
GRAY_BG   = RGBColor(0xF2, 0xF2, 0xF2)
WHITE     = RGBColor(0xFF, 0xFF, 0xFF)
TH_BLUE   = RGBColor(0x2E, 0x54, 0x96)


def set_cell_bg(cell, color: RGBColor):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    hex_color = '{:02X}{:02X}{:02X}'.format(color[0], color[1], color[2])
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def set_para_bg(para, color: RGBColor):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    hex_color = '{:02X}{:02X}{:02X}'.format(color[0], color[1], color[2])
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    pPr.append(shd)


def add_heading1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = DARK_BLUE
    # Alt çizgi
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '8')
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), '1F3864')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def add_heading2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = MED_BLUE
    return p


def add_heading3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = MED_BLUE
    return p


def add_para(doc, text, bold_parts=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    if bold_parts:
        parts = text.split('**')
        for i, part in enumerate(parts):
            run = p.add_run(part)
            run.font.size = Pt(11)
            run.font.color.rgb = BLACK
            if i % 2 == 1:
                run.bold = True
    else:
        run = p.add_run(text)
        run.font.size = Pt(11)
        run.font.color.rgb = BLACK
    return p


def add_meta(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.color.rgb = GRAY_TXT
    return p


def add_terminal(doc, lines):
    """lines: list of (text, style) where style is 'cmd', 'out', or 'ok'"""
    for text, style in lines:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Cm(0.5)
        set_para_bg(p, GRAY_BG)
        run = p.add_run(text)
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
        if style == 'cmd':
            run.font.color.rgb = GREEN
            run.bold = True
        elif style == 'ok':
            run.font.color.rgb = GREEN
        else:
            run.font.color.rgb = BLACK
    # Boşluk sonrası
    doc.add_paragraph().paragraph_format.space_after = Pt(6)


def add_note(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.right_indent = Cm(0.5)
    set_para_bg(p, YELLOW_BG)
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.color.rgb = BLACK


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    # Başlık satırı
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        hdr_cells[i].paragraphs[0].runs[0].bold = True
        hdr_cells[i].paragraphs[0].runs[0].font.color.rgb = WHITE
        hdr_cells[i].paragraphs[0].runs[0].font.size = Pt(10)
        set_cell_bg(hdr_cells[i], TH_BLUE)
    # Veri satırları
    for r_idx, row in enumerate(rows):
        cells = table.rows[r_idx + 1].cells
        for c_idx, val in enumerate(row):
            cells[c_idx].text = val
            cells[c_idx].paragraphs[0].runs[0].font.size = Pt(10)
            if r_idx % 2 == 1:
                set_cell_bg(cells[c_idx], RGBColor(0xF7, 0xF7, 0xF7))
    doc.add_paragraph().paragraph_format.space_after = Pt(6)


def add_divider(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'CCCCCC')
    pBdr.append(bottom)
    pPr.append(pBdr)


# ─────────────────────────────────────────────
# DÖKÜMAN İÇERİĞİ
# ─────────────────────────────────────────────

add_heading1(doc, 'Qwen3.6-35B-A3B-FP8 Model Devreye Alım Dökümanı')
add_meta(doc, 'Tarih: 21 Nisan 2026   |   Hazırlayan: TAI – Yapay Zeka Ekibi   |   Sunucu: plusaigpu01')

# 1
add_heading2(doc, '1. Amaç ve Kapsam')
add_para(doc, 'Bu döküman, Qwen3.6-35B-A3B-FP8 büyük dil modelinin **plusaigpu01** sunucusundaki GPU 7 (NVIDIA H200 144G) üzerine kurulumu, yapılandırılması ve servise alınması sürecini adım adım açıklamaktadır.', bold_parts=True)
add_para(doc, 'Model, vLLM OpenAI-uyumlu API sunucusu üzerinden **8015 portu** ile hizmete açılmış; mevcut chatboss Docker ağına dahil edilmiştir.', bold_parts=True)

# 2
add_heading2(doc, '2. Sistem Bilgileri')
add_table(doc,
    ['Parametre', 'Değer'],
    [
        ['Sunucu', 'plusaigpu01'],
        ['GPU', 'NVIDIA H200 SXM 141GiB × 8 adet'],
        ['Kullanılan GPU', 'GPU 7 (~70 GB boş VRAM)'],
        ['Model', 'Qwen/Qwen3.6-35B-A3B-FP8'],
        ['Model Boyutu', '34.90 GiB (42 safetensor parçası)'],
        ['Mimari', 'Mixture of Experts (MoE) – 35B toplam, 3B aktif parametre, FP8'],
        ['Çalışma Portu', '8015'],
        ['Model Dizini', '/data02/models/Qwen3.6-35B-A3B-FP8'],
        ['Compose Dizini', '/data01/kurulum/qwen3-35b/'],
        ['Docker İmajı', 'vllm/vllm-openai:latest (v0.19.1)'],
        ['Docker Ağı', 'chatboss (external)'],
    ]
)

add_divider(doc)

# 3
add_heading2(doc, '3. GPU Durumu Kontrolü')
add_para(doc, 'Kurulum öncesinde sunucudaki GPU\'ların anlık durumu ve kullanım bilgisi kontrol edilmiştir.')
add_terminal(doc, [
    ('$ nvidia-smi', 'cmd'),
    ('+-----------------------------------------------------------------------------------------+', 'out'),
    ('| NVIDIA-SMI 570.133.07   Driver Version: 570.133.07   CUDA Version: 12.8               |', 'out'),
    ('|=========================================================================================|', 'out'),
    ('|   0  NVIDIA H200 SXM 141GiB   On   00000000:31:00.0   Off |   N/A                     |', 'out'),
    ('|   1  NVIDIA H200 SXM 141GiB   On   00000000:35:00.0   Off |   N/A                     |', 'out'),
    ('|   2  NVIDIA H200 SXM 141GiB   On   00000000:3D:00.0   Off |   N/A                     |', 'out'),
    ('|   3  NVIDIA H200 SXM 141GiB   On   00000000:41:00.0   Off |   N/A                     |', 'out'),
    ('|   4  NVIDIA H200 SXM 141GiB   On   00000000:B1:00.0   Off |   N/A                     |', 'out'),
    ('|   5  NVIDIA H200 SXM 141GiB   On   00000000:B5:00.0   Off |   N/A                     |', 'out'),
    ('|   6  NVIDIA H200 SXM 141GiB   On   00000000:BD:00.0   Off |   N/A                     |', 'out'),
    ('|   7  NVIDIA H200 SXM 141GiB   On   00000000:C1:00.0   Off |   N/A                     |', 'out'),
    ('+-----------------------------------------------------------------------------------------+', 'out'),
])
add_note(doc, 'Not: GPU 7, diğer modeller tarafından kullanılmamakta olup ~71 GB boş VRAM ile Qwen3.6-35B-A3B-FP8 modeli için en uygun seçenektir.')

add_divider(doc)

# 4
add_heading2(doc, '4. Model İndirme')
add_heading3(doc, '4.1 Hedef Dizin Oluşturma')
add_terminal(doc, [
    ('$ mkdir -p /data02/models/Qwen3.6-35B-A3B-FP8', 'cmd'),
])

add_heading3(doc, '4.2 Git LFS ile Model İndirme')
add_para(doc, 'Kurumsal ağda PyPI erişimi engellendiğinden huggingface-cli kullanılamamış, bunun yerine Git LFS ile doğrudan HuggingFace\'den indirme yapılmıştır.')
add_terminal(doc, [
    ('$ git lfs install', 'cmd'),
    ('Git LFS initialized.', 'out'),
    ('', 'out'),
    ('$ cd /data02/models', 'cmd'),
    ('$ git clone https://huggingface.co/Qwen/Qwen3.6-35B-A3B-FP8 Qwen3.6-35B-A3B-FP8', 'cmd'),
    ('Cloning into \'Qwen3.6-35B-A3B-FP8\'...', 'out'),
    ('remote: Enumerating objects: 86, done.', 'out'),
    ('remote: Counting objects: 100% (86/86), done.', 'out'),
    ('remote: Compressing objects: 100% (85/85), done.', 'out'),
    ('Receiving objects: 100% (86/86), done.', 'out'),
    ('', 'out'),
    ('Filtering content: 100% (42/42), 34.90 GiB | 269.00 MiB/s, done.', 'out'),
])
add_note(doc, 'Not: İndirme işlemi ~269 MB/s hızla gerçekleşmiş ve toplam 34.90 GiB boyutundaki model başarıyla indirilmiştir.')

add_heading3(doc, '4.3 Model Dosyalarının Doğrulanması')
add_terminal(doc, [
    ('$ ls -lh /data02/models/Qwen3.6-35B-A3B-FP8/', 'cmd'),
    ('total 34G', 'out'),
    ('-rw-r--r-- 1 root root  663 config.json', 'out'),
    ('-rw-r--r-- 1 root root  242 generation_config.json', 'out'),
    ('-rw-r--r-- 1 root root 841M model-00001-of-00042.safetensors', 'out'),
    ('-rw-r--r-- 1 root root 980M model-00002-of-00042.safetensors', 'out'),
    ('...', 'out'),
    ('-rw-r--r-- 1 root root 980M model-00042-of-00042.safetensors', 'out'),
    ('-rw-r--r-- 1 root root  88K model.safetensors.index.json', 'out'),
    ('-rw-r--r-- 1 root root 7.0M tokenizer.json', 'out'),
    ('-rw-r--r-- 1 root root 2.4K tokenizer_config.json', 'out'),
])

add_divider(doc)

# 5
add_heading2(doc, '5. Docker Compose Yapılandırması')
add_heading3(doc, '5.1 Dizin Yapısının Hazırlanması')
add_terminal(doc, [
    ('$ mkdir -p /data01/kurulum/qwen3-35b', 'cmd'),
    ('$ cd /data01/kurulum/qwen3-35b', 'cmd'),
])

add_heading3(doc, '5.2 docker-compose.yml Dosyası')
add_para(doc, 'Model GPU 7\'ye sabitlenmiş, port 8015 olarak atanmış ve gpu-memory-utilization değeri GPU\'daki mevcut kullanıma göre 0.48 olarak ayarlanmıştır.')
add_terminal(doc, [
    ('name: qwen3-35b', 'out'),
    ('', 'out'),
    ('networks:', 'out'),
    ('  chatboss:', 'out'),
    ('    external: true', 'out'),
    ('', 'out'),
    ('x-deploy-7: &gpu-7-deploy', 'out'),
    ('  deploy:', 'out'),
    ('    resources:', 'out'),
    ('      reservations:', 'out'),
    ('        devices:', 'out'),
    ("          - driver: nvidia", 'out'),
    ("            device_ids: [ '7' ]", 'out'),
    ('            capabilities: [ gpu ]', 'out'),
    ('', 'out'),
    ('services:', 'out'),
    ('  qwen3-35b:', 'out'),
    ('    image: vllm/vllm-openai:latest', 'out'),
    ('    volumes:', 'out'),
    ('      - /data02/models/Qwen3.6-35B-A3B-FP8:/models/Qwen3.6-35B-A3B-FP8:ro', 'out'),
    ('    command: >', 'out'),
    ('      --model /models/Qwen3.6-35B-A3B-FP8', 'out'),
    ('      --served-model-name Qwen3.6-35B-A3B', 'out'),
    ('      --gpu-memory-utilization 0.48', 'out'),
    ('      --tensor-parallel-size 1', 'out'),
    ('      --max-model-len 32768', 'out'),
    ('      --host 0.0.0.0', 'out'),
    ('      --port 8000', 'out'),
    ('    ports:', 'out'),
    ('      - "8015:8000"', 'out'),
    ('    <<: *gpu-7-deploy', 'out'),
    ('    ipc: host', 'out'),
    ('    networks:', 'out'),
    ('      - chatboss', 'out'),
])
add_note(doc, 'gpu-memory-utilization: 0.48 — GPU 7\'de başka işlemler ~70 GB VRAM kullandığından kalan alan için bu oran belirlenmiştir. KV önbelleği için 28.33 GiB ayrılmış, 370.656 token ve yaklaşık 40x eşzamanlı istek kapasitesi sağlanmıştır.')

add_divider(doc)

# 6
add_heading2(doc, '6. vLLM Docker İmajının Güncellenmesi')
add_para(doc, 'Qwen3.6-35B-A3B-FP8 modeli eski vLLM sürümleriyle (v0.11.0) desteklenmemektedir. Gerekli mimari (Qwen3_5MoeForConditionalGeneration) için vLLM\'in son sürümü indirilmiştir.')
add_terminal(doc, [
    ('$ docker pull vllm/vllm-openai:latest', 'cmd'),
    ('latest: Pulling from vllm/vllm-openai', 'out'),
    ('...', 'out'),
    ('Status: Downloaded newer image for vllm/vllm-openai:latest', 'ok'),
    ('docker.io/library/vllm/vllm-openai:latest', 'out'),
])

add_divider(doc)

# 7
add_heading2(doc, '7. Servisin Başlatılması')
add_heading3(doc, '7.1 Konteyner Başlatma')
add_terminal(doc, [
    ('$ cd /data01/kurulum/qwen3-35b', 'cmd'),
    ('$ docker compose up -d', 'cmd'),
    ('[+] Running 1/1', 'out'),
    ('  Container qwen3-35b-qwen3-35b-1  Started', 'ok'),
])

add_heading3(doc, '7.2 Başlatma Logları')
add_terminal(doc, [
    ('$ docker compose logs -f', 'cmd'),
    ('INFO:     Loading model weights...', 'out'),
    ('INFO:     GPU memory allocated: 34.26 GiB', 'out'),
    ('INFO:     KV cache allocated: 28.33 GiB (370,656 tokens)', 'out'),
    ('INFO:     Max concurrency: ~40x', 'out'),
    ('INFO:     Application startup complete.', 'ok'),
    ('INFO:     Uvicorn running on http://0.0.0.0:8000 (Press CTRL+C to quit)', 'out'),
])
add_note(doc, 'Başlatma süresi: Model yüklemesi ve KV önbelleği hazırlığı yaklaşık 3–5 dakika sürmektedir. Logda "Application startup complete." mesajı görünene kadar beklenmesi gerekmektedir.')

add_divider(doc)

# 8
add_heading2(doc, '8. Doğrulama Testleri')
add_heading3(doc, '8.1 Model Listesi Kontrolü')
add_terminal(doc, [
    ('$ curl -s http://localhost:8015/v1/models | python3 -m json.tool', 'cmd'),
    ('{', 'out'),
    ('    "object": "list",', 'out'),
    ('    "data": [', 'out'),
    ('        {', 'out'),
    ('            "id": "Qwen3.6-35B-A3B",', 'out'),
    ('            "object": "model",', 'out'),
    ('            "owned_by": "vllm",', 'out'),
    ('            "max_model_len": 32768', 'out'),
    ('        }', 'out'),
    ('    ]', 'out'),
    ('}', 'out'),
])

add_heading3(doc, '8.2 Sohbet Testi')
add_terminal(doc, [
    ('$ curl -s http://localhost:8015/v1/chat/completions \\', 'cmd'),
    ('  -H "Content-Type: application/json" \\', 'cmd'),
    ('  -d \'{"model":"Qwen3.6-35B-A3B","messages":[{"role":"user",\'', 'cmd'),
    ('     \'"content":"BMW X serisi hakkında bilgi verir misin? /no_think"}]}\'', 'cmd'),
    ('{', 'out'),
    ('    "choices": [{', 'out'),
    ('        "message": {', 'out'),
    ('            "role": "assistant",', 'out'),
    ('            "content": "BMW X serisi, Bavyera Motorlu Araçlar İşleri (BMW)', 'out'),
    ('             tarafından üretilen SAV segmentindeki araç ailesidir..."', 'out'),
    ('        }', 'out'),
    ('    }]', 'out'),
    ('}', 'out'),
])
add_note(doc, 'Önemli – Thinking Modu: Model varsayılan olarak yanıt öncesinde İngilizce <think>...</think> blokları üretmektedir. Bu blokları devre dışı bırakmak için istek metnine /no_think eklenmesi yeterlidir.')

add_divider(doc)

# 9
add_heading2(doc, '9. Port ve Ağ Durumu')
add_terminal(doc, [
    ('$ ss -tlnp | grep 8015', 'cmd'),
    ('LISTEN  0  4096  0.0.0.0:8015  0.0.0.0:*  users:(("docker-proxy",...))', 'out'),
    ('', 'out'),
    ('$ docker ps --filter "name=qwen3-35b"', 'cmd'),
    ('CONTAINER ID   IMAGE                    STATUS        PORTS', 'out'),
    ('a3f812c1b490   vllm/vllm-openai:latest  Up 2 hours    0.0.0.0:8015->8000/tcp', 'out'),
])

add_divider(doc)

# 10
add_heading2(doc, '10. Model ve Kaynak Kullanımı Özeti')
add_table(doc,
    ['Kaynak', 'Değer'],
    [
        ['GPU VRAM – Model', '~34.26 GiB'],
        ['GPU VRAM – KV Önbellek', '28.33 GiB'],
        ['Toplam GPU VRAM Tahsisi', '~62 GiB / 141 GiB'],
        ['KV Token Kapasitesi', '370.656 token'],
        ['Eşzamanlı İstek Kapasitesi', '~40 istek'],
        ['Maksimum Bağlam Uzunluğu', '32.768 token'],
        ['API Endpoint', 'http://plusaigpu01:8015/v1'],
        ['Model Adı (API)', 'Qwen3.6-35B-A3B'],
    ]
)

add_divider(doc)

# 11
add_heading2(doc, '11. Dosya ve Dizin Yapısı')
add_terminal(doc, [
    ('/data01/kurulum/qwen3-35b/', 'out'),
    ('└── docker-compose.yml', 'out'),
    ('', 'out'),
    ('/data02/models/Qwen3.6-35B-A3B-FP8/', 'out'),
    ('├── config.json', 'out'),
    ('├── generation_config.json', 'out'),
    ('├── model-00001-of-00042.safetensors', 'out'),
    ('│   ...', 'out'),
    ('├── model-00042-of-00042.safetensors', 'out'),
    ('├── model.safetensors.index.json', 'out'),
    ('└── tokenizer.json', 'out'),
])

add_divider(doc)

# 12
add_heading2(doc, '12. Karşılaşılan Sorunlar ve Çözümleri')
add_table(doc,
    ['Sorun', 'Çözüm'],
    [
        ['Root disk doluluk hatası (/data tam)', 'Model /data02/models altına yönlendirildi'],
        ['PyPI erişimi engelli (kurumsal proxy)', 'pip yerine git lfs kullanıldı'],
        ['vLLM v0.11.0 Qwen3 mimarisini desteklemiyor', 'vllm/vllm-openai:latest (v0.19.1) çekildi'],
        ['gpu-memory-utilization 0.85 OOM hatası verdi', 'Değer 0.48\'e düşürüldü'],
        ['Yanıtlar İngilizce <think> bloğu içeriyor', '/no_think komutu ile düşünce modu kapatılıyor'],
    ]
)

add_divider(doc)

# 13
add_heading2(doc, '13. Kullanım Notları')
add_para(doc, 'Servis yeniden başlatmak gerektiğinde aşağıdaki komutlar kullanılmalıdır:')
add_terminal(doc, [
    ('$ cd /data01/kurulum/qwen3-35b', 'cmd'),
    ('$ docker compose down', 'cmd'),
    ('$ docker compose up -d', 'cmd'),
    ('$ docker compose logs -f', 'cmd'),
])
add_para(doc, 'API üzerinden modeli sorgulamak için OpenAI uyumlu herhangi bir istemci kullanılabilir. Base URL: **http://plusaigpu01:8015/v1** — Model adı: **Qwen3.6-35B-A3B**', bold_parts=True)
add_note(doc, 'Thinking modu hakkında: Model varsayılan olarak çıktı öncesinde İngilizce düşünce bloğu (<think>...</think>) üretmektedir. Arayüzde gösterilmek istenmiyorsa /no_think sistem mesajı olarak eklenmeli ya da API yanıtında bu bloklar filtrelenmelidir.')

# Kaydet
output_path = '/home/user/mlops-platform/docs/qwen3-35b-devreye-alim.docx'
doc.save(output_path)
print(f'Dosya oluşturuldu: {output_path}')
